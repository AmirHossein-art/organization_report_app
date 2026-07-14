from io import BytesIO

import pandas as pd

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from utils.format_helpers import (
    report_type_label,
    to_jalali_date,
    to_jalali_datetime,
    to_persian_digits,
)


def build_dashboard_excel(
    period: dict,
    rows: list[dict],
) -> bytes:
    output = BytesIO()

    total_rows = len(rows)
    submitted_rows = [row for row in rows if row["has_report"]]
    missing_rows = [row for row in rows if not row["has_report"]]
    late_rows = [row for row in rows if row["is_late"]]
    on_time_rows = [
        row for row in rows
        if row["has_report"] and not row["is_late"]
    ]
    with_files_rows = [
        row for row in rows
        if row["file_count"] > 0
    ]

    participation_rate = 0

    if total_rows > 0:
        participation_rate = round((len(submitted_rows) / total_rows) * 100, 1)

    summary_df = pd.DataFrame(
        [
            {
                "شاخص": "عنوان بازه",
                "مقدار": period["title"],
            },
            {
                "شاخص": "نوع گزارش",
                "مقدار": report_type_label(period["report_type"]),
            },
            {
                "شاخص": "شروع بازه",
                "مقدار": to_jalali_date(period["period_start"]),
            },
            {
                "شاخص": "پایان بازه",
                "مقدار": to_jalali_date(period["period_end"]),
            },
            {
                "شاخص": "تعداد مورد انتظار",
                "مقدار": to_persian_digits(total_rows),
            },
            {
                "شاخص": "تعداد ثبت‌شده",
                "مقدار": to_persian_digits(len(submitted_rows)),
            },
            {
                "شاخص": "تعداد ثبت‌نشده",
                "مقدار": to_persian_digits(len(missing_rows)),
            },
            {
                "شاخص": "تعداد به‌موقع",
                "مقدار": to_persian_digits(len(on_time_rows)),
            },
            {
                "شاخص": "تعداد تأخیری",
                "مقدار": to_persian_digits(len(late_rows)),
            },
            {
                "شاخص": "تعداد دارای پیوست",
                "مقدار": to_persian_digits(len(with_files_rows)),
            },
            {
                "شاخص": "درصد مشارکت",
                "مقدار": f"{to_persian_digits(participation_rate)}٪",
            },
        ]
    )

    overview_df = pd.DataFrame(
        [
            {
                "کاربر": row["user_full_name"],
                "نام کاربری": row["username"],
                "پروژه": row["project_title"],
                "بازه": row["period_title"],
                "وضعیت": row["status_label"],
                "تاریخ ثبت": to_jalali_datetime(row["submitted_at"]) if row["submitted_at"] else "-",
                "تعداد پیوست": to_persian_digits(row["file_count"]),
                "فایل‌های پیوست": "، ".join([file["original_filename"] for file in row.get("files", [])]),
            }
            for row in rows
        ]
    )

    submitted_df = pd.DataFrame(
        [
            {
                "کاربر": row["user_full_name"],
                "نام کاربری": row["username"],
                "پروژه": row["project_title"],
                "وضعیت": row["status_label"],
                "تاریخ ثبت": to_jalali_datetime(row["submitted_at"]) if row["submitted_at"] else "-",
                "تعداد پیوست": to_persian_digits(row["file_count"]),
                "فایل‌های پیوست": "، ".join([file["original_filename"] for file in row.get("files", [])]),
                "فعالیت‌های انجام‌شده": row["activities_done"],
                "نتایج حاصل‌شده": row["results_achieved"],
                "اقدامات آتی": row["next_actions"],
                "شاخص‌ها": row["kpi_text"],
            }
            for row in submitted_rows
        ]
    )

    missing_df = pd.DataFrame(
        [
            {
                "کاربر": row["user_full_name"],
                "نام کاربری": row["username"],
                "پروژه": row["project_title"],
                "بازه": row["period_title"],
                "وضعیت": row["status_label"],
            }
            for row in missing_rows
        ]
    )

    late_df = pd.DataFrame(
        [
            {
                "کاربر": row["user_full_name"],
                "نام کاربری": row["username"],
                "پروژه": row["project_title"],
                "تاریخ ثبت": to_jalali_datetime(row["submitted_at"]) if row["submitted_at"] else "-",
                "تعداد پیوست": to_persian_digits(row["file_count"]),
                "فایل‌های پیوست": "، ".join([file["original_filename"] for file in row.get("files", [])]),
                "فعالیت‌های انجام‌شده": row["activities_done"],
                "نتایج حاصل‌شده": row["results_achieved"],
                "اقدامات آتی": row["next_actions"],
                "شاخص‌ها": row["kpi_text"],
            }
            for row in late_rows
        ]
    )

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        summary_df.to_excel(writer, sheet_name="Summary", index=False)
        overview_df.to_excel(writer, sheet_name="Overview", index=False)
        submitted_df.to_excel(writer, sheet_name="Submitted", index=False)
        missing_df.to_excel(writer, sheet_name="Missing", index=False)
        late_df.to_excel(writer, sheet_name="Late", index=False)

        workbook = writer.book

        for worksheet in workbook.worksheets:
            worksheet.sheet_view.rightToLeft = True

            for column_cells in worksheet.columns:
                max_length = 0
                column_letter = column_cells[0].column_letter

                for cell in column_cells:
                    if cell.value is not None:
                        max_length = max(max_length, len(str(cell.value)))

                adjusted_width = min(max_length + 4, 60)
                worksheet.column_dimensions[column_letter].width = adjusted_width

    output.seek(0)

    return output.getvalue()


def build_dashboard_word(
    period: dict,
    summary: dict,
    rows: list[dict],
    ai_summary: str = None,
) -> bytes:
    doc = Document()
    
    # تنظیم استایل پیش‌فرض برای پشتیبانی بهتر از فارسی
    style = doc.styles['Normal']
    style.font.name = 'Vazirmatn'
    style.font.size = Pt(11)

    # عنوان اصلی
    heading = doc.add_heading(f"گزارش مدیریتی بازه: {period['title']}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"نوع گزارش: {report_type_label(period['report_type'])}")
    doc.add_paragraph(f"تاریخ شروع: {to_jalali_date(period['period_start'])}")
    doc.add_paragraph(f"تاریخ پایان: {to_jalali_date(period['period_end'])}")
    doc.add_paragraph(f"درصد مشارکت: {to_persian_digits(summary.get('participation_rate', 0))}٪")
    
    if ai_summary:
        doc.add_heading("تحلیل و خلاصه مدیریتی (تولید شده توسط هوش مصنوعی)", level=2)
        p_ai = doc.add_paragraph(ai_summary)
        p_ai.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

    # فیلتر کردن داده‌ها
    submitted_rows = [row for row in rows if row["has_report"] and not row["is_late"]]
    late_rows = [row for row in rows if row["is_late"]]
    missing_rows = [row for row in rows if not row["has_report"]]

    def add_report_section(title, report_list):
        if not report_list:
            return
            
        doc.add_heading(title, level=2)
        for row in report_list:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f"کاربر: {row['user_full_name']} | پروژه: {row['project_title']}").bold = True
            
            submit_date = to_jalali_datetime(row["submitted_at"]) if row["submitted_at"] else "-"
            doc.add_paragraph(f"تاریخ ثبت: {submit_date}")
            
            doc.add_paragraph("فعالیت‌های انجام‌شده:").bold = True
            doc.add_paragraph(row.get("activities_done") or "ثبت نشده")
            
            doc.add_paragraph("نتایج حاصل‌شده:").bold = True
            doc.add_paragraph(row.get("results_achieved") or "ثبت نشده")
            
            doc.add_paragraph("اقدامات آتی:").bold = True
            doc.add_paragraph(row.get("next_actions") or "ثبت نشده")
            
            doc.add_paragraph("شاخص‌ها:").bold = True
            doc.add_paragraph(row.get("kpi_text") or "ثبت نشده")
            
            # فایل‌های پیوست
            files = row.get("files", [])
            if files:
                doc.add_paragraph("فایل‌های پیوست:").bold = True
                for file in files:
                    doc.add_paragraph(f"- {file['original_filename']}")
                    
            doc.add_paragraph("-" * 40)

    # افزودن بخش‌های مختلف به فایل
    add_report_section("گزارش‌های ثبت‌شده (به‌موقع)", submitted_rows)
    add_report_section("گزارش‌های تأخیری", late_rows)

    # افزودن لیست ثبت‌نشده‌ها
    if missing_rows:
        doc.add_heading("کاربران / پروژه‌های بدون گزارش", level=2)
        for row in missing_rows:
            p = doc.add_paragraph(f"کاربر: {row['user_full_name']} | پروژه: {row['project_title']}")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ذخیره در حافظه و بازگرداندن بایت‌ها
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()